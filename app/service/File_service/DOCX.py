import os 
import uuid
from docx import Document
from fastapi import HTTPException
from langchain_core.documents import Document as LCDocument 

from .base import FileProcessor
from .FileFactory import FileProcessorRegistry

@FileProcessorRegistry.register(".docx")
class DOCXProcessor(FileProcessor):
    def __init__(self):
        self.upload_dir = "data/upload/docx"
        self.indexes_dir = "data/indexes/docx"
        os.makedirs(self.upload_dir,exist_ok=True)
        os.makedirs(self.indexes_dir,exist_ok = True)

    def save_file(self,file_bytes, file_name):
        file_id = uuid.uuid4().hex #.uuid4() create random UUID, .hex() convert it into 32-character string 
        safe_name = os.path.basename(file_name)
        new_filename = f"{file_id}_{safe_name}"
        file_path = os.path.join(self.upload_dir,new_filename)
        with open(file_path,"wb") as f:
            f.write(file_bytes)
        return file_id 
    
    def get_file_path(self,folder,file_id):
        base_dir = self.upload_dir if folder == "upload" else self.indexes_dir
        files = [f for f in os.listdir(base_dir) if f.startswith(f"{file_id}_")]
        if not files:
            return os.path.join(base_dir, file_id)
        return os.path.join(base_dir, files[0]) 
    
    def get_file(self,file_id,upload_file_path=None):
        upload_file_path = (self.get_file_path("upload",file_id) if upload_file_path is None else upload_file_path)
        return Document(upload_file_path) #docx.document.Document 
    
    def process_file(self,file_id,upload_file_path=None):  
        try:
            doc = self.get_file(file_id,upload_file_path)
            raw_name = (os.path.basename(self.get_file_path("upload",file_id)) if upload_file_path is None else os.path.basename(upload_file_path))
            file_name = raw_name.replace(f"{file_id}_", "", 1)
            texts = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n".join(texts)
            return [LCDocument(page_content=full_text,
                              metadata={"title":doc.core_properties.title,
                                        "file_name":file_name,
                                        "file_id": file_id})] #return List[Langchain Document]
        except HTTPException:
            raise
        except Exception:
            raise HTTPException(status_code=500,detail="Failed to process DOCX file")
        
    def is_complicated_file(self,doc): #docx.document.Document
        # 1. Examine the existence of table 
        if len(doc.tables) > 0:
            for table in doc.tables:
                # Check for nested tables (Count the number of sub-boards in each cell.)
                for row in table.rows:
                    for cell in row.cells:
                        if len(cell.tables) > 0:
                            return True # Nested Tables
                
                # Check Merged cells
                # In python-docx, if actual cells differs from rows*columns
                if len(table._cells) != (len(table.rows) * len(table.columns)):
                    return True # Merged Cells

        # 2. Check Textbox and Floating Objects
        # They're usually in 'inline_shapes' or 'shapes'
        if len(doc.inline_shapes) > 3:
            return True # Too much pictures or embedding textbox 
        return False 
    
    def get_list_file(self):
        return [tuple(f.split("_", 1)) for f in os.listdir(self.upload_dir)]